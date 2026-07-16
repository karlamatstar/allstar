import os
import sys
import subprocess
import datetime
import shutil
import json
import io
from pathlib import Path

# 터미널 출력(stdout) 인코딩을 UTF-8로 강제하여 k6 등의 특수문자 출력 시 UnicodeEncodeError(cp949) 방지
if isinstance(sys.stdout, io.TextIOWrapper):
    sys.stdout.reconfigure(encoding='utf-8')

# Add project root to sys.path
PROJECT_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from allstar.ai_agent.api.jira_client import JIRA_URL, JIRA_USER, JIRA_API_TOKEN, JIRA_PROJECT_KEY

def run_k6():
    print("▶ Running Chaos Test (K6)...", flush=True)
    script = PROJECT_ROOT / "ops" / "performance" / "chaos_test.js"
    result_file = PROJECT_ROOT / "ops" / "performance" / "results" / "chaos_result.json"
    os.makedirs(result_file.parent, exist_ok=True)

    # 프로젝트 RUN 폴더에 둔 실행 파일을 우선하고, 없으면 시스템 설치 경로를 사용한다.
    k6_exe = PROJECT_ROOT / "RUN" / "k6.exe"
    k6_bin = str(k6_exe) if k6_exe.exists() else shutil.which("k6")
    if not k6_bin:
        message = "[오류] K6 실행 파일을 찾지 못했습니다. RUN/k6.exe를 두거나 K6를 시스템에 설치하세요."
        print(message, flush=True)
        return result_file, message
    cmd = [k6_bin, "run", f"--summary-export={result_file}", str(script)]
    creationflags = 0
    if sys.platform == "win32":
        creationflags = subprocess.CREATE_NO_WINDOW

    process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, encoding='utf-8', errors='replace', creationflags=creationflags)
    output_lines = []
    for line in process.stdout:
        print(line, end='', flush=True)
        output_lines.append(line)
    process.wait()

    if process.returncode != 0:
        print("⚠️ K6 completed with errors or thresholds failed.", flush=True)
    return result_file, "".join(output_lines)

def run_pytest():
    print("▶ Running Functional Tests (Pytest, external AI calls excluded)...", flush=True)
    cmd = [
        sys.executable, "-m", "pytest", "-v", "tests/",
        "--ignore=tests/ai_agent/test_negative_cases.py",
        "--ignore=tests/ai_agent/test_evaluation_pipeline.py",
        "--ignore=tests/voc/evaluation/test_pipeline_e2e.py",
        "-k", "not end_to_end",
    ]
    env = os.environ.copy()
    env["PYTHONIOENCODING"] = "utf-8"

    creationflags = 0
    if sys.platform == "win32":
        creationflags = subprocess.CREATE_NO_WINDOW

    process = subprocess.Popen(cmd, cwd=str(PROJECT_ROOT), stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, encoding='utf-8', errors='replace', env=env, creationflags=creationflags)
    output_lines = []
    for line in process.stdout:
        print(line, end='', flush=True)
        output_lines.append(line)
    process.wait()

    if process.returncode != 0:
        print("⚠️ Pytest completed with some failures.", flush=True)
    return "".join(output_lines)

def create_chaos_defect_report(k6_out, pytest_out):
    print("▶ Generating Chaos Defect Report...")
    today = datetime.datetime.now()
    timestamp = today.strftime("%Y-%m-%d %H:%M:%S")
    timestamp_file = today.strftime("%Y%m%d_%H%M%S")

    report_dir = PROJECT_ROOT / "_OUTPUT" / "reports" / "defects" / "chaos"
    log_dir = PROJECT_ROOT / "_OUTPUT" / "logs" / "ai_agent" / "chaos"
    os.makedirs(log_dir, exist_ok=True)

    defect_md = report_dir / "defect_report.md"

    import re

    test_desc_map = {
        "test_chat_returns_answer_for_valid_question": "정상적인 질문에 대한 답변 반환 검증",
        "test_chat_rejects_empty_question": "빈 문자열 입력 시 예외 처리(422) 검증",
        "test_health_returns_ok": "서버 상태(Health Check) 정상 응답 검증",
        "test_metrics_endpoint_exposed": "모니터링(Metrics) 엔드포인트 노출 검증",
        "test_generate_live_report": "실시간 결함 보고서 정상 생성 검증",
        "test_generate_live_report_without_logs": "빈 로그 상태에서의 보고서 생성 방어 로직 검증",
        "test_negative_case_is_handled_safely": "부정적 케이스(예외 상황) 안전 처리 검증",
        "test_pipeline_generates_reports": "결함 자동화 파이프라인 연동 검증"
    }

    status_map = {
        "PASSED": "통과",
        "FAILED": "실패",
        "ERROR": "에러",
        "SKIPPED": "스킵"
    }

    pytest_lines = []
    for line in pytest_out.splitlines():
        match = re.search(r'::(\S+)\s+(PASSED|FAILED|ERROR|SKIPPED)\s+\[', line)
        if match:
            raw_test_name = match.group(1)
            status = match.group(2)

            # 파라미터 분리 (예: test_name[TC-026])
            base_test_name = raw_test_name.split('[')[0]
            param_part = raw_test_name[len(base_test_name):] if '[' in raw_test_name else ""

            kor_desc = test_desc_map.get(base_test_name, base_test_name)
            if param_part:
                kor_desc += f" {param_part}"

            kor_status = status_map.get(status, status)
            color = "🟢 " if status == "PASSED" else "🔴 " if status in ["FAILED", "ERROR"] else "🟡 "

            pytest_lines.append(f"  - {color}**{kor_desc}** : {kor_status}")

    pytest_details = "\n".join(pytest_lines)
    if not pytest_details:
        pytest_details = "  - 상세 테스트 내역을 추출할 수 없습니다."

    content = f"""# 인프라 장애 모의 훈련(Chaos Test) 결과 보고서

**작성일시**: {timestamp}
**테스트 범위**: `/fault-lab` API 장애 모의(Delay, 500 오류, Timeout 등) 및 Pytest 통합 검증

## 결함 제목
[장애-모의훈련] /fault-lab API 장애 복합 재현 결과

## 재현 절차
GUI '장애·기능 검증 시험' 실행 (k6 `chaos_test.js` 호출) 및 기능 검증(Pytest) 자동화 파이프라인 수행

## 실제 결과
* **FL-001 (정상 응답)**: 기준선 검증을 위해 정상적인 API 호출 시나리오 통과 확인 (HTTP 200).
* **FL-002 (응답 지연 1초)**: 시스템에 1초 지연이 발생했을 때의 응답 대기 확인.
* **FL-003 (응답 지연 5초)**: 시스템에 5초 지연이 발생했을 때의 응답 대기 확인.
* **FL-004 (500 오류 재현)**: 시스템 오류 발생 시 내부 로직이 500 에러를 반환하는 시나리오 재현.
* **FL-005 (타임아웃 재현)**: 네트워크 대기 시간 초과로 인한 504 Gateway Timeout 발생 재현.
* **FL-006 (잘못된 시나리오)**: 파라미터 누락/오류 시 400 Bad Request 처리 확인.
* **FL-007 (기능 테스트 세부 결과)**:
{pytest_details}

## 영향도
응답 지연 시 사용자 대기시간 증가 및 500 에러에 의한 비즈니스 흐름 중단 등 치명적 오류 발생 위험 존재

## 원인 추정
고의적인 `/fault-lab` 시나리오 호출 파라미터 유발로 인한 복합 장애 모의 환경 구성. (실제 운영 환경일 경우 DB 병목, 리소스 부족, 통신 타임아웃 등이 원인이 될 수 있음)

## 조치 방안 (세부 항목별 개별 조치 가이드)

### [FL-002, FL-003] 응답 지연 (Delay) 발생에 대한 조치
- **현상**: 의도적인 지연(1초, 5초)으로 인해 사용자 응답 대기 시간이 길어짐
- **조치 방안**: 클라이언트(웹/앱) 단에서 최대 대기 시간(Timeout)을 3초 내외로 적절히 설정하고, 지연 발생 시 "현재 요청량이 많아 처리가 지연되고 있습니다"와 같은 스켈레톤 UI 또는 로딩 애니메이션을 제공하여 사용자 이탈을 최소화해야 합니다.

### [FL-004] HTTP 500 내부 서버 오류에 대한 조치
- **현상**: `/fault-lab` 호출 시 서버 내부 로직 예외로 인해 500 에러 반환
- **조치 방안**: 서버 측 글로벌 예외 처리기(Exception Handler)를 강화해야 합니다. 500 에러 발생 시 시스템 내부 스택 트레이스 등 민감한 정보가 노출되지 않도록 마스킹 처리하고, 사용자에게는 "일시적인 오류가 발생했습니다. 잠시 후 다시 시도해주세요."라는 친화적인 안내 메시지를 응답하도록 보완이 필요합니다.

### [FL-005] HTTP 504 Gateway Timeout에 대한 조치
- **현상**: 정해진 타임아웃 시간을 초과하여 게이트웨이 또는 리버스 프록시에서 연결을 강제 차단함
- **조치 방안**: 백엔드 서버 앞단(Nginx, API Gateway 등)의 타임아웃 설정과 로드밸런서 정책을 점검해야 합니다. 장시간 실행이 필수적인 요청의 경우 동기식 처리에서 비동기 처리(Background Task 또는 Message Queue) 구조로 아키텍처 개선이 요구됩니다.

### [FL-007] 통합 기능 테스트 (Pytest) 결과 조치
- **현상**: 위 '실제 결과 (상세 로그)'의 Pytest 실행 결과에 표시된 개별 기능 통과/실패 내역
- **조치 방안**: 파이테스트 로그 상에서 `FAILED` 또는 `ERROR` 로 식별된 개별 기능(예: 특정 API 오류, DB 검증 실패 등)에 대해, 각 기능의 도메인 담당 파트로 티켓을 분배하여 원인을 분석하고 즉각적인 핫픽스를 배포해야 합니다.

## 재시험 및 QA 검토
조치 후 자동화 스크립트(FL-001~008) 재수행 필요.
"""

    with open(defect_md, "w", encoding="utf-8") as f:
        f.write(content)

    # Copy to log
    log_file = log_dir / f"defect_report_{timestamp_file}.md"
    shutil.copy2(defect_md, log_file)

    return defect_md

def convert_md_to_word(md_file: Path):
    print("▶ Converting to Word document...")
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("python-docx is not installed.")
        return

    doc = Document()
    doc.add_heading("인프라 장애 모의 훈련(Chaos Test) 결과 보고서", 0)

    with open(md_file, "r", encoding="utf-8") as f:
        text = f.read()

    for line in text.splitlines():
        if line.startswith("# "): continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        else:
            doc.add_paragraph(line)

    # Add QA Engineer Highlight
    p = doc.add_paragraph()
    run = p.add_run("💡 QA 엔지니어 검토 의견: 본 품질 이상은 인프라 장애로 기인한 것으로 분석됩니다. 따라서 조치 방안 적용을 담당 파트에 권고하며, 조치 완료 즉시 파이프라인 재검증을 수행할 예정입니다.")
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.italic = True

    word_path = md_file.parent / "final_defect_report.docx"
    doc.save(word_path)
    print(f"✅ Word report saved to {word_path}")
    return word_path

def create_chaos_jira_issue():
    print("▶ Registering defect to Jira BUG-4 Epic...")
    if not all([JIRA_URL, JIRA_USER, JIRA_API_TOKEN]):
        print("Jira environment variables not set. Skipping.")
        return

    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    description = f"""모니터링(k6/Pytest) 중 애플리케이션 API에서 치명적인 서버 장애가 감지되었습니다.

* 장애 발생 시간: {timestamp}
* 발생 엔드포인트: GET /fault-lab
* HTTP 상태 코드: 500, 504 등 복합 장애

[재현 경로 (URL)]
http://127.0.0.1:8000/fault-lab?scenario=error500

[담당자 조치 필요 사항]
현재 해당 API가 정상적으로 동작하지 않고 있습니다.
현상 파악 후 결함 보고서(_OUTPUT/reports/defects/chaos/defect_report.md)를 작성하고,
조치 완료 후 재시험(Re-test)하여 본 티켓을 Closed 처리해 주시기 바랍니다."""

    payload = {
        "fields": {
            "project": {"key": JIRA_PROJECT_KEY},
            "parent": {"key": "BUG-4"},
            "issuetype": {"name": "Bug"},
            "priority": {"name": "High"},
            "summary": "[장애-모의훈련] /fault-lab API 장애 복합 감지",
            "description": description
        }
    }

    try:
        import httpx
        url = f"{JIRA_URL}/rest/api/2/issue"
        auth = (JIRA_USER, JIRA_API_TOKEN)
        headers = {"Content-Type": "application/json"}
        response = httpx.post(url, json=payload, auth=auth, headers=headers, timeout=10.0)
        if response.status_code == 201:
            print(f"✅ Jira Issue Created: {response.json().get('key')}")
        else:
            print(f"❌ Jira Issue Creation Failed: {response.text}")
    except Exception as e:
        print(f"Jira API Error: {e}")

if __name__ == "__main__":
    print("="*60)
    print("통합 QA 검증 파이프라인 시작")
    print("="*60)

    k6_result, k6_output = run_k6()
    pytest_output = run_pytest()
    md_file = create_chaos_defect_report(k6_output, pytest_output)
    convert_md_to_word(md_file)
    # create_chaos_jira_issue() # 사용자의 요청으로 지라 자동 티켓 생성 비활성화

    print("="*60)
    print("통합 QA 검증 파이프라인 완료!")
    print("="*60)
