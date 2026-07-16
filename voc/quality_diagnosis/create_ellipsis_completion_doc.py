"""기존 보고서를 수정하지 않고 말줄임된 문장의 완성본만 별도 DOCX로 만든다."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ROOT.parent
REPORTS = PROJECT_ROOT / "quality" / "reports" / "voc" / "testcase"
CSV_PATH = REPORTS / "llm_judge_result.csv"
OUTPUT_PATH = REPORTS / "보고서_말줄임문장_완성본.docx"


SECTIONS = [
    (
        "VOC_개선_QA_점검완료보고서.docx",
        [
            ("즉시 보류 사례 표", ["TC-08", "TC-15"]),
        ],
    ),
    (
        "VOC_분석_종합품질평가보고서.docx",
        [
            ("4장 우수 사례와 취약 사례 - 상위 5개 사례 표", ["TC-10", "TC-12", "TC-06", "TC-04", "TC-01"]),
            ("4장 우수 사례와 취약 사례 - 하위 5개 사례 표", ["TC-13", "TC-14", "TC-08", "TC-15", "TC-16"]),
            ("5장 즉시 보류와 데이터 한계 - 즉시 보류 사례 표", ["TC-08", "TC-15"]),
        ],
    ),
]


def load_rationales() -> dict[str, str]:
    with CSV_PATH.open("r", encoding="utf-8-sig", newline="") as file:
        rows = list(csv.DictReader(file))
    result = {}
    for row in rows:
        case_id = (row.get("case_id") or "").strip()
        rationale = (row.get("rationale") or "").strip()
        if case_id and rationale:
            result[case_id] = rationale
    return result


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10)

    for style_name, size, color in (
        ("Title", 24, "17365D"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 12, "2F5597"),
    ):
        style = doc.styles[style_name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_case_table(doc: Document, case_ids: list[str], rationales: dict[str, str]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(2.1)
    table.columns[1].width = Cm(14.8)

    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ("케이스", "교체해서 넣을 완성 문장")):
        set_cell_shading(cell, "D9EAF7")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True

    for case_id in case_ids:
        if case_id not in rationales:
            raise KeyError(f"CSV에서 {case_id} rationale을 찾을 수 없습니다")
        cells = table.add_row().cells
        cells[0].vertical_alignment = 1
        label = cells[0].paragraphs[0]
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = label.add_run(case_id)
        run.bold = True
        paragraph = cells[1].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.add_run(rationales[case_id])

    doc.add_paragraph()


def build_document() -> Path:
    rationales = load_rationales()
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("보고서 말줄임 문장 완성본")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("기존 DOCX 원본은 수정하지 않고, 중간에 끊긴 문장의 교체용 전체 문장만 수록")
    run.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89)

    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    set_cell_shading(note.cell(0, 0), "FFF2CC")
    note.cell(0, 0).paragraphs[0].add_run(
        "사용 방법: 아래에서 원본 문서와 표 위치를 확인한 뒤, 해당 TC의 말줄임 문장을 "
        "이 문서의 완성 문장으로 교체합니다. 같은 TC가 여러 표에 있으면 각 위치에 동일하게 적용합니다."
    )
    doc.add_paragraph()

    item_number = 0
    for source_name, groups in SECTIONS:
        doc.add_heading(source_name, level=1)
        for location, case_ids in groups:
            doc.add_heading(location, level=2)
            item_number += len(case_ids)
            add_case_table(doc, case_ids, rationales)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"교체 문장 {item_number}개 · 원본 문서 변경 없음")

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
